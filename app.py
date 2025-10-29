from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
from datetime import datetime

# Only import modules that work on Vercel
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# OCR modules (won't work on Vercel)
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

app = Flask(__name__)
CORS(app)

# Configuration
# Check if running on Vercel
IS_VERCEL = os.environ.get('VERCEL') == '1'

if OCR_AVAILABLE and not IS_VERCEL:
    # For local development
    if os.path.exists(r'C:\Program Files\Tesseract-OCR\tesseract.exe'):
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    # For Linux/other platforms
    elif os.path.exists('/usr/bin/tesseract'):
        pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

UPLOAD_FOLDER = '/tmp/uploads' if os.path.exists('/tmp') else 'uploads'
OUTPUT_FOLDER = '/tmp/output' if os.path.exists('/tmp') else 'output'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'pdf', 'docx', 'csv', 'xlsx', 'xls'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 4.5 * 1024 * 1024  # 4.5MB for Vercel

# Create necessary directories (only if not on Vercel at build time)
try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
except:
    pass

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_image(image_path, lang='eng+khm'):
    """Process a single image and extract text"""
    if not OCR_AVAILABLE:
        raise Exception("OCR is not available on this platform. Please use Railway.app, Render.com, or Heroku for full OCR support.")
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=lang)
        return text
    except Exception as e:
        raise Exception(f"Error processing image: {str(e)}")

def process_pdf(pdf_path, lang='eng+khm'):
    """Process PDF and extract text from all pages"""
    if not OCR_AVAILABLE:
        raise Exception("PDF OCR is not available on this platform. Please use Railway.app, Render.com, or Heroku for full OCR support.")
    try:
        pages = convert_from_path(pdf_path, dpi=300)
        all_text = ""
        
        for i, page in enumerate(pages):
            text = pytesseract.image_to_string(page, lang=lang)
            all_text += f"\n\n--- Page {i+1} ---\n{text}"
        
        return all_text
    except Exception as e:
        raise Exception(f"Error processing PDF: {str(e)}")

def process_docx(docx_path):
    """Process DOCX file and extract text"""
    if not DOCX_AVAILABLE:
        raise Exception("DOCX processing is not available. Missing python-docx module.")
    try:
        doc = Document(docx_path)
        all_text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            all_text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            all_text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                all_text += row_text + "\n"
        
        return all_text.strip()
    except Exception as e:
        raise Exception(f"Error processing DOCX: {str(e)}")

def process_csv(csv_path):
    """Process CSV file and extract text"""
    if not PANDAS_AVAILABLE:
        raise Exception("CSV processing is not available. Missing pandas module.")
    try:
        df = pd.read_csv(csv_path)
        
        # Convert dataframe to formatted text
        all_text = "--- CSV Data ---\n\n"
        all_text += df.to_string(index=False)
        
        return all_text
    except Exception as e:
        raise Exception(f"Error processing CSV: {str(e)}")

def process_xlsx(xlsx_path):
    """Process XLSX file and extract text from all sheets"""
    if not PANDAS_AVAILABLE or not OPENPYXL_AVAILABLE:
        raise Exception("XLSX processing is not available. Missing pandas or openpyxl module.")
    try:
        all_text = ""
        
        # Load the workbook
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
        
        # Process each sheet
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            all_text += f"\n\n--- Sheet: {sheet_name} ---\n\n"
            
            # Read data using pandas for better formatting
            df = pd.read_excel(xlsx_path, sheet_name=sheet_name)
            all_text += df.to_string(index=False)
        
        return all_text.strip()
    except Exception as e:
        raise Exception(f"Error processing XLSX: {str(e)}")

@app.route('/')
def index():
    """Serve the main UI"""
    return render_template('index.html')

@app.route('/api/health', methods=['GET'])
def health():
    """Health check endpoint"""
    features = {
        'ocr': OCR_AVAILABLE,
        'docx': DOCX_AVAILABLE,
        'csv': PANDAS_AVAILABLE,
        'xlsx': OPENPYXL_AVAILABLE,
        'platform': 'vercel' if IS_VERCEL else 'other'
    }
    return jsonify({
        'status': 'healthy', 
        'message': 'Khmer OCR API is running',
        'features': features
    })

@app.route('/api/ocr', methods=['POST'])
def ocr():
    """Main OCR endpoint"""
    # Show helpful error for Vercel users
    if IS_VERCEL:
        file_ext = request.files.get('file', type(''))
        if hasattr(file_ext, 'filename') and file_ext.filename:
            ext = file_ext.filename.rsplit('.', 1)[-1].lower()
            if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'pdf']:
                return jsonify({
                    'error': 'OCR is not supported on Vercel. For full OCR functionality, please deploy to Railway.app, Render.com, or Heroku. See RAILWAY_DEPLOY.md for instructions.',
                    'docs_url': 'https://github.com/PhonSobon/khmerOCR_Tesseract/blob/main/RAILWAY_DEPLOY.md'
                }), 501
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Allowed types: PNG, JPG, JPEG, GIF, BMP, TIFF, PDF, DOCX, CSV, XLSX'}), 400
    
    try:
        # Get language preference
        lang = request.form.get('lang', 'eng+khm')
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(filepath)
        
        # Process based on file type
        file_ext = filename.rsplit('.', 1)[1].lower()
        
        if file_ext == 'pdf':
            extracted_text = process_pdf(filepath, lang)
        elif file_ext == 'docx':
            extracted_text = process_docx(filepath)
        elif file_ext == 'csv':
            extracted_text = process_csv(filepath)
        elif file_ext in ['xlsx', 'xls']:
            extracted_text = process_xlsx(filepath)
        else:
            extracted_text = process_image(filepath, lang)
        
        # Save output text
        output_filename = f"{timestamp}_output.txt"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        
        return jsonify({
            'success': True,
            'text': extracted_text,
            'filename': filename,
            'output_file': output_filename
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
    finally:
        # Clean up uploaded file
        try:
            if 'filepath' in locals() and os.path.exists(filepath):
                os.remove(filepath)
        except:
            pass

@app.route('/api/download/<filename>', methods=['GET'])
def download(filename):
    """Download extracted text file"""
    try:
        filepath = os.path.join(app.config['OUTPUT_FOLDER'], secure_filename(filename))
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("Starting Khmer OCR API Server...")
    print("Access the UI at: http://localhost:5000")
    app.run(debug=True, host='0.0.0.0', port=5000)

# For Vercel
app = app
